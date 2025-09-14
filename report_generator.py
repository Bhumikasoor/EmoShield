import matplotlib.pyplot as plt
from collections import Counter
from docx import Document
from docx.shared import Inches
import io

def generate_doc_report(results):
    doc = Document()
    doc.add_heading("EmoShield Chat Analysis Report", level=0)

    # ---- Summary ----
    total = len(results)
    high = sum(1 for r in results if r["risk_level"] == "HIGH")
    medium = sum(1 for r in results if r["risk_level"] == "MEDIUM")
    low = sum(1 for r in results if r["risk_level"] == "LOW")

    doc.add_paragraph(f"Total messages analyzed: {total}")
    doc.add_paragraph(f"HIGH risk: {high}")
    doc.add_paragraph(f"MEDIUM risk: {medium}")
    doc.add_paragraph(f"LOW risk: {low}")

    # ---- Risk Distribution Pie Chart ----
    counts = Counter([r["risk_level"] for r in results])
    labels = list(counts.keys())
    sizes = list(counts.values())

    fig, ax = plt.subplots(figsize=(4,4))
    ax.pie(sizes, labels=labels, autopct="%1.1f%%", colors=["red","orange","green"])
    ax.set_title("Risk Level Distribution")

    img_stream = io.BytesIO()
    fig.savefig(img_stream, format="png")
    plt.close(fig)
    img_stream.seek(0)
    doc.add_heading("Risk Level Distribution", level=1)
    doc.add_picture(img_stream, width=Inches(4))

    # ---- Sentiment Distribution Bar Chart ----
    sentiments = Counter([r["sentiment"] for r in results])
    labels = list(sentiments.keys())
    values = list(sentiments.values())

    fig, ax = plt.subplots(figsize=(5,4))
    ax.bar(labels, values, color="blue")
    ax.set_title("Sentiment Distribution")
    ax.set_xlabel("Sentiment")
    ax.set_ylabel("Count")

    img_stream2 = io.BytesIO()
    fig.savefig(img_stream2, format="png")
    plt.close(fig)
    img_stream2.seek(0)
    doc.add_heading("Sentiment Distribution", level=1)
    doc.add_picture(img_stream2, width=Inches(5))

    # ---- Top Keywords ----
    all_keywords = []
    for r in results:
        if r["keywords"] != "None":
            all_keywords.extend(r["keywords"].split(", "))
    keyword_counts = Counter(all_keywords).most_common(5)

    if keyword_counts:
        doc.add_heading("Top Keywords", level=1)
        for kw, cnt in keyword_counts:
            doc.add_paragraph(f"{kw}: {cnt}")

    # ---- Detailed Table ----
    doc.add_heading("Detailed Results", level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Line"
    hdr_cells[1].text = "Message"
    hdr_cells[2].text = "Risk"
    hdr_cells[3].text = "Score"
    hdr_cells[4].text = "Keywords"

    for r in results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r["line"])
        row_cells[1].text = r["message"][:40] + ("..." if len(r["message"]) > 40 else "")
        row_cells[2].text = r["risk_level"]
        row_cells[3].text = str(r["total_score"])
        row_cells[4].text = r["keywords"]

    # ---- Save to BytesIO instead of disk ----
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer
